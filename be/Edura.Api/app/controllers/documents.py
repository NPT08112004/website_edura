# app/controllers/documents.py
# -*- coding: utf-8 -*-

from flask import Blueprint, request, jsonify, current_app, Response, stream_with_context
from werkzeug.utils import secure_filename
from bson.objectid import ObjectId
from bson.errors import InvalidId
from jwt import ExpiredSignatureError, InvalidTokenError

import os
import re
import uuid
import tempfile
import shutil
import subprocess
import platform
from io import BytesIO
from collections import Counter
import concurrent.futures
import textwrap
import requests
import jwt  # pyjwt
import fitz  # PyMuPDF
from PIL import Image, ImageDraw, ImageFont
import boto3
from urllib.parse import urlparse
from datetime import datetime, timedelta, date

from app.services.aws_service import aws_service
from app.services.ai_service import ai_service
from app.services.mongo_service import mongo_collections
from app.models.document import Document
from jwt import ExpiredSignatureError, InvalidTokenError
from flask import current_app
from bson import ObjectId

documents_bp = Blueprint("documents", __name__, url_prefix="/api/documents")

DOC_COST = int(os.getenv("DOC_COST_POINTS", "5"))  # mặc định 5 điểm/tài liệu

# Cho phép định dạng
ALLOWED_DOC_EXT = {"pdf", "docx", "doc"}
ALLOWED_IMG_EXT = {"png", "jpg", "jpeg", "webp"}

# ====== Cờ tối ưu qua ENV (A-tweaks) ======
SKIP_WORD_CONVERSION = os.getenv("SKIP_WORD_CONVERSION", "false").lower() == "true"
USE_AI = os.getenv("USE_AI", "true").lower() == "true"


# ===================== Helpers =====================

def _allowed_ext(filename: str, allow: set[str]) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in allow



def _naive_keywords(text: str, k: int = 12) -> list[str]:
    words = re.findall(r"[a-zA-ZÀ-ỹ0-9]{3,}", (text or "").lower())
    stop = {"the","and","for","with","that","this","from","have","you","are","not","your","of","to","in","on","by","is",
            "các","của","và","cho","những","một","được","trên","trong","khi","này","đó","là","đến","từ","có"}
    words = [w for w in words if w not in stop]
    return [w for w,_ in Counter(words).most_common(k)]


def _extract_text_from_pdf_bytes(file_bytes: bytes, max_pages: int = 6) -> str:
    """Trích text PDF (PyMuPDF) – giảm còn 6 trang đầu để tăng tốc."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    texts = []
    for i in range(min(max_pages, doc.page_count)):
        texts.append(doc.load_page(i).get_text("text"))
    return "\n".join(texts).strip()


def _extract_text_from_pdf_bytes_smart(file_bytes: bytes, max_pages: int = 50) -> str:
    """
    Trích text PDF từ 50 trang đầu tiên cho tài liệu dài.
    Với tài liệu quá dài (> 100 trang): chỉ lấy 50 trang đầu để tóm tắt.
    Thử nhiều phương pháp để đảm bảo trích được text.
    """
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        total_pages = doc.page_count
        texts = []
        
        # Chỉ lấy max_pages trang đầu tiên (mặc định 50 trang)
        pages_to_extract = min(max_pages, total_pages)
        print(f"[Extract Text] Trích text từ {pages_to_extract}/{total_pages} trang đầu tiên")
        
        # Phương pháp 1: Thử get_text("text") - phương pháp chuẩn
        for i in range(pages_to_extract):
            try:
                page = doc.load_page(i)
                # Thử nhiều phương pháp
                page_text = page.get_text("text")
                if not page_text or not page_text.strip():
                    # Thử phương pháp khác
                    page_text = page.get_text("dict")
                    if page_text and "blocks" in page_text:
                        text_parts = []
                        for block in page_text["blocks"]:
                            if "lines" in block:
                                for line in block["lines"]:
                                    if "spans" in line:
                                        for span in line["spans"]:
                                            if "text" in span:
                                                text_parts.append(span["text"])
                        page_text = " ".join(text_parts)
                    else:
                        page_text = ""
                
                if page_text and page_text.strip():
                    texts.append(page_text.strip())
            except Exception as e:
                print(f"[Extract Text] Lỗi khi trích text từ trang {i}: {e}")
                continue
        
        result = "\n".join(texts).strip()
        print(f"[Extract Text] Phương pháp 1 (PyMuPDF): {len(result)} ký tự từ {len(texts)} trang")
        
        # Nếu không có text, thử phương pháp 2: pdfplumber
        if not result or len(result) < 100:
            print(f"[Extract Text] Thử phương pháp 2: pdfplumber...")
            try:
                import pdfplumber
                from io import BytesIO
                texts_plumber = []
                with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                    for i, page in enumerate(pdf.pages[:pages_to_extract]):
                        try:
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                texts_plumber.append(page_text.strip())
                        except Exception as e:
                            print(f"[Extract Text] pdfplumber lỗi trang {i}: {e}")
                            continue
                
                result_plumber = "\n".join(texts_plumber).strip()
                print(f"[Extract Text] Phương pháp 2 (pdfplumber): {len(result_plumber)} ký tự từ {len(texts_plumber)} trang")
                
                if len(result_plumber) > len(result):
                    result = result_plumber
            except ImportError:
                print(f"[Extract Text] pdfplumber không có sẵn, bỏ qua")
            except Exception as e:
                print(f"[Extract Text] Lỗi khi dùng pdfplumber: {e}")
        
        doc.close()
        
        if result:
            print(f"[Extract Text] Trích text thành công: {len(result)} ký tự")
        else:
            print(f"[Extract Text] CẢNH BÁO: Không trích được text từ PDF (có thể là PDF scan)")
        
        return result
    except Exception as e:
        print(f"[Extract Text] Lỗi khi mở PDF: {e}")
        import traceback
        traceback.print_exc()
        return ""


def _ocr_text_from_pdf_bytes(file_bytes: bytes, pages_max: int = 5, scale: float = 2.0) -> str:
    """OCR fallback nếu PDF là scan – chỉ dùng khi thật sự cần."""
    try:
        import pytesseract
    except ImportError:
        print(f"[OCR] pytesseract không có sẵn, bỏ qua OCR")
        return ""
    except Exception as e:
        print(f"[OCR] Lỗi import pytesseract: {e}")
        return ""
    
    try:
        print(f"[OCR] Bắt đầu OCR {pages_max} trang đầu tiên...")
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        texts = []
        pages_to_ocr = min(pages_max, doc.page_count)
        
        for i in range(pages_to_ocr):
            try:
                page = doc.load_page(i)
                pix = page.get_pixmap(alpha=False, matrix=fitz.Matrix(scale, scale))
                if pix.alpha:
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                t = pytesseract.image_to_string(img, lang="eng+vie")
                if t and t.strip():
                    texts.append(t.strip())
                    print(f"[OCR] Trang {i+1}: {len(t)} ký tự")
            except Exception as e:
                print(f"[OCR] Lỗi OCR trang {i}: {e}")
                continue
        
        result = "\n".join(texts).strip()
        print(f"[OCR] OCR thành công: {len(result)} ký tự từ {len(texts)} trang")
        doc.close()
        return result
    except Exception as e:
        print(f"[OCR] Lỗi khi OCR: {e}")
        import traceback
        traceback.print_exc()
        return ""


def _generate_thumb_from_pdf_bytes(file_bytes: bytes, scale: float = 1.0) -> BytesIO | None:
    """Render trang đầu PDF -> JPEG thumbnail (qua PIL)."""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        if doc.page_count == 0:
            return None
        pix = doc.load_page(0).get_pixmap(alpha=False, matrix=fitz.Matrix(scale, scale))
        if pix.alpha:
            pix = fitz.Pixmap(fitz.csRGB, pix)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        buf = BytesIO()
        img.save(buf, format="JPEG", optimize=True, quality=85)
        buf.seek(0)
        return buf
    except Exception:
        return None


def _get_pdf_page_count(file_bytes: bytes) -> int:
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        return int(doc.page_count or 0)
    except Exception:
        return 0


def _infer_document_page_count(document: dict) -> int:
    try:
        url = document.get("s3_url") or document.get("s3Url")
        if not url:
            return 0
        ext = url.split("?", 1)[0].rsplit(".", 1)[-1].lower()
        response = requests.get(url, timeout=45)
        if response.status_code >= 400:
            return 0
        content = response.content
        if ext == "pdf":
            return _get_pdf_page_count(content)
        if ext in {"docx", "doc"}:
            converted = _convert_word_to_pdf_bytes(content, ext)
            if converted:
                return _get_pdf_page_count(converted)
        return 0
    except Exception:
        return 0


def _generate_placeholder_thumb(title: str, width: int = 800, height: int = 500) -> BytesIO:
    """Placeholder thumbnail khi không có PDF để render."""
    bg, fg, accent = (245, 248, 255), (33, 37, 41), (0, 123, 255)
    img = Image.new("RGB", (width, height), bg)
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle([20, 20, width-20, height-20], radius=24, outline=accent, width=3)
    try:
        font_big = ImageFont.truetype("arial.ttf", 32)
        font_small = ImageFont.truetype("arial.ttf", 20)
    except Exception:
        font_big = ImageFont.load_default()
        font_small = ImageFont.load_default()
    draw.text((40, 40), "Document Preview", fill=accent, font=font_small)
    draw.multiline_text((40, 90), textwrap.fill(title or "Tài liệu", width=28), fill=fg, font=font_big, spacing=8)
    buf = BytesIO()
    img.save(buf, format="JPEG", optimize=True, quality=85)
    buf.seek(0)
    return buf


def _convert_word_to_pdf_bytes(file_bytes: bytes, ext: str) -> bytes | None:
    """Convert .doc/.docx -> PDF bằng LibreOffice; fallback docx2pdf."""
    ext = ext.lower().lstrip(".")
    if ext not in ("docx", "doc"):
        return None
    tmpdir = tempfile.mkdtemp(prefix="conv_")
    src_path = os.path.join(tmpdir, f"input.{ext}")
    try:
        with open(src_path, "wb") as f:
            f.write(file_bytes)

        soffice = shutil.which("soffice")
        if not soffice and platform.system() == "Windows":
            for c in [r"C:\Program Files\LibreOffice\program\soffice.exe",
                      r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"]:
                if os.path.isfile(c):
                    soffice = c
                    break

        if soffice:
            try:
                subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir", tmpdir, src_path],
                               check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                pdf_path = os.path.join(tmpdir, "input.pdf")
                if os.path.exists(pdf_path):
                    return open(pdf_path, "rb").read()
            except subprocess.CalledProcessError:
                pass

        if ext == "docx":
            try:
                from docx2pdf import convert
                convert(src_path, tmpdir)
                pdf_path = os.path.join(tmpdir, "input.pdf")
                if os.path.exists(pdf_path):
                    return open(pdf_path, "rb").read()
            except Exception:
                pass
        return None
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


def _extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    """Trích text trực tiếp từ DOCX (docx2txt -> python-docx)."""
    tmpdir = tempfile.mkdtemp(prefix="docx_")
    path = os.path.join(tmpdir, "input.docx")
    try:
        with open(path, "wb") as f:
            f.write(file_bytes)
        try:
            import docx2txt
            return (docx2txt.process(path) or "").strip()
        except Exception:
            pass
        try:
            import docx
            d = docx.Document(path)
            return "\n".join([p.text for p in d.paragraphs if p.text]).strip()
        except Exception:
            return ""
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


def _ensure_lookup_ids(school_id: str, category_id: str) -> tuple[str, str]:
    """Đảm bảo có sẵn school/category mặc định để tránh lỗi rỗng."""
    if not school_id:
        s = mongo_collections.schools.find_one({}, {"_id": 1})
        if not s:
            mongo_collections.schools.insert_many([
                {"name": "ĐH Bách Khoa TP.HCM"},
                {"name": "ĐH Khoa học Tự nhiên TP.HCM"},
                {"name": "ĐH CNTT (UIT)"},
            ])
            s = mongo_collections.schools.find_one({}, {"_id": 1})
        school_id = str(s["_id"])
    if not category_id:
        c = mongo_collections.categories.find_one({}, {"_id": 1})
        if not c:
            mongo_collections.categories.insert_many([
                {"name": "Toán cao cấp"},
                {"name": "Cấu trúc dữ liệu & Giải thuật"},
                {"name": "Cơ sở dữ liệu"},
                {"name": "Mạng máy tính"},
                {"name": "Hệ điều hành"},
                {"name": "Kinh tế vi mô"},
                {"name": "Marketing căn bản"},
            ])
            c = mongo_collections.categories.find_one({}, {"_id": 1})
        category_id = str(c["_id"])
    return school_id, category_id


def _apply_rate_limit_if_available(route_func):
    """Apply rate limiting nếu flask-limiter có sẵn"""
    try:
        from flask import current_app
        limiter = current_app.config.get('LIMITER')
        if limiter:
            return limiter.limit("10 per minute")(route_func)
    except Exception:
        pass
    return route_func

def _get_current_user_strict(return_doc: bool = False):
    """Đọc Bearer token & trả (ObjectId user, tên hiển thị)."""
    auth = request.headers.get("Authorization", "")
    if not auth.lower().startswith("bearer "):
        raise InvalidTokenError("Thiếu Bearer token")
    token = auth.split(" ", 1)[1].strip()
    # Bắt buộc JWT_KEY từ config hoặc env, không có fallback
    secret = (current_app.config.get("JWT_KEY")
              or os.getenv("JWT_KEY"))
    if not secret:
        raise ValueError("JWT_KEY chưa được cấu hình trong environment variables")
    payload = jwt.decode(token, secret, algorithms=["HS256"], options={"verify_aud": False})
    uid_raw = payload.get("userId") or payload.get("id") or payload.get("_id") or payload.get("sub")
    uid = ObjectId(str(uid_raw))
    u = mongo_collections.users.find_one({"_id": uid}, {"username": 1, "fullName": 1, "name": 1, "email": 1, "avatarUrl": 1})
    if not u:
        raise InvalidTokenError("User không tồn tại")
    uploader_name = u.get("fullName") or u.get("username") or u.get("name") or u.get("email")
    if return_doc:
        return uid, u
    return uid, uploader_name


def _get_current_user_optional():
    auth = request.headers.get("Authorization", "")
    if not auth.lower().startswith("bearer "):
        return None, None
    token = auth.split(" ", 1)[1].strip()
    secret = (
        current_app.config.get("JWT_KEY")
        or os.getenv("JWT_KEY")
        or current_app.config.get("JWT_SECRET")
        or current_app.config.get("JWT_SECRET_KEY")
        or "dev_secret"
    )
    try:
        payload = jwt.decode(token, secret, algorithms=["HS256"], options={"verify_aud": False})
        uid_raw = payload.get("userId") or payload.get("id") or payload.get("_id") or payload.get("sub")
        uid = ObjectId(str(uid_raw))
        u = mongo_collections.users.find_one({"_id": uid}, {"username": 1, "fullName": 1, "avatarUrl": 1})
        if not u:
            return None, None
        return uid, u
    except Exception:
        return None, None


# ===================== NEW: Direct-to-S3 Presign =====================

@documents_bp.route("/presign", methods=["POST"])
def presign_upload():
    """
    Cấp URL ký sẵn để FE upload trực tiếp lên S3 (PUT).
    Body: { "ext": "pdf|docx|doc", "contentType": "..." }
    Trả: { "key": "documents/<uuid>.<ext>", "url": "<presigned_put_url>" }
    """
    data = request.get_json() or {}
    ext = (data.get("ext") or "pdf").lower().strip(".")
    if ext not in ALLOWED_DOC_EXT:
        return jsonify({"error": "Định dạng không hợp lệ"}), 400

    ct = data.get("contentType") or "application/octet-stream"
    # xác thực để ràng buộc quyền cấp presign
    try:
        _uid, _ = _get_current_user_strict()
    except Exception as e:
        return jsonify({"error": f"Auth lỗi: {e}"}), 401

    bucket = os.getenv("S3_BUCKET_NAME")
    region = os.getenv("AWS_REGION", "ap-southeast-1")
    key = f"documents/{uuid.uuid4()}.{ext}"

    s3 = boto3.client("s3", region_name=region,
                      aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
                      aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY"))
    url = s3.generate_presigned_url(
        ClientMethod="put_object",
        Params={"Bucket": bucket, "Key": key, "ContentType": ct},
        ExpiresIn=900  # 15 phút
    )
    return jsonify({"key": key, "url": url}), 200


@documents_bp.route("/register", methods=["POST"])
@_apply_rate_limit_if_available
def register_document():
    """
    Đăng ký metadata sau khi FE đã PUT file lên S3.
    Body: { title, schoolId, categoryId, s3Key, imageUrl? }
    Trả: { document_id }
    Xử lý AI/thumbnail chạy nền (non-blocking).
    """
    try:
        data = request.get_json() or {}
        title = (data.get("title") or "").strip()
        s3_key = (data.get("s3Key") or "").strip()
        image_url = (data.get("imageUrl") or "").strip() or None
        school_id = (data.get("schoolId") or "").strip()
        category_id = (data.get("categoryId") or "").strip()

        # Validate title
        try:
            from app.utils.validation import validate_title
            title = validate_title(title)
        except ValueError as e:
            return jsonify({"error": str(e)}), 400

        if not s3_key:
            return jsonify({"error": "Thiếu s3Key"}), 400

        school_id, category_id = _ensure_lookup_ids(school_id, category_id)

        # xác thực user
        try:
            current_user_oid, uploader_name = _get_current_user_strict()
        except ExpiredSignatureError:
            return jsonify({"error": "JWT hết hạn. Vui lòng đăng nhập lại."}), 401
        except InvalidTokenError as e:
            return jsonify({"error": f"Token không hợp lệ: {e}"}), 401

        bucket = os.getenv("S3_BUCKET_NAME")
        region = os.getenv("AWS_REGION", "ap-southeast-1")
        s3_url = f"https://{bucket}.s3.{region}.amazonaws.com/{s3_key}"

        # Tạo document (đặt summary/keywords tạm)
        doc = Document(
            title=title,
            s3_url=s3_url,
            user_id=str(current_user_oid),
            summary=f"Đang xử lý tóm tắt cho: {title}",
            keywords=["processing"],
            school_id=str(school_id),
            category_id=str(category_id),
            image_url=image_url,
            pages=None,
        )
        doc_dict = doc.to_mongo_doc()
        doc_dict["uploaderName"] = uploader_name
        result = mongo_collections.documents.insert_one(doc_dict)
        doc_id = result.inserted_id

        # Xử lý AI/thumbnail bất đồng bộ
        def _bg_enrich():
            try:
                r = requests.get(s3_url, timeout=45)
                if r.status_code >= 400:
                    return
                pdf_bytes = r.content if s3_key.lower().endswith(".pdf") else None
                if not pdf_bytes and s3_key.lower().endswith((".docx", ".doc")):
                    pdf_bytes = _convert_word_to_pdf_bytes(r.content, s3_key.rsplit(".", 1)[-1])

                text = ""
                if pdf_bytes:
                    # Sử dụng hàm smart để trích text từ nhiều phần cho tài liệu dài
                    page_count = _get_pdf_page_count(pdf_bytes)
                    if page_count > 100:
                        # Tài liệu quá dài (>100 trang): chỉ lấy 50 trang đầu để tóm tắt
                        text = _extract_text_from_pdf_bytes_smart(pdf_bytes, max_pages=50)
                    else:
                        # Tài liệu ngắn: lấy tất cả hoặc 6 trang đầu
                        text = _extract_text_from_pdf_bytes(pdf_bytes, max_pages=min(6, page_count))
                    if not text or len(text.strip()) < 100:
                        print(f"[Register] Text quá ngắn ({len(text) if text else 0} ký tự), thử OCR...")
                        # Với tài liệu dài, OCR nhiều trang hơn để có đủ nội dung
                        ocr_pages = 10 if page_count > 100 else 5
                        ocr_text = _ocr_text_from_pdf_bytes(pdf_bytes, pages_max=ocr_pages, scale=2.0)
                        if ocr_text and len(ocr_text.strip()) > len(text.strip() if text else ""):
                            text = ocr_text
                            print(f"[Register] OCR thành công, sử dụng text từ OCR: {len(text)} ký tự")
                        elif not text:
                            text = ocr_text  # Sử dụng OCR text ngay cả khi ngắn nếu không có text nào
                        
                        # Nếu vẫn không có text (PDF scan, tesseract không có), thử Gemini Vision API
                        if (not text or len(text.strip()) < 50) and USE_AI and ai_service:
                            print(f"[Register] Thử sử dụng Gemini Vision API để OCR và tóm tắt trực tiếp...")
                            try:
                                vision_summary, vision_keywords = ai_service.extract_and_summarize_from_pdf_images(
                                    pdf_bytes, max_pages=10, page_count=page_count
                                )
                                if vision_summary:
                                    # Sử dụng summary từ Gemini Vision, không cần text nữa
                                    summary = vision_summary
                                    keywords = vision_keywords
                                    print(f"[Register] Gemini Vision thành công: summary={len(vision_summary)} ký tự")
                                    # Đặt text rỗng vì đã có summary từ Vision API
                                    text = ""  # Không cần text nữa vì đã có summary
                            except Exception as e:
                                print(f"[Register] Lỗi khi dùng Gemini Vision: {e}")
                                import traceback
                                traceback.print_exc()
                elif s3_key.lower().endswith(".docx"):
                    # không convert tại đây để tiết kiệm thời gian; chỉ fallback text thô nếu cần
                    pass

                # Lưu ý: summary có thể đã được tạo từ Gemini Vision API ở trên
                if 'summary' not in locals():
                    summary, keywords = None, None
                # Chỉ tóm tắt bằng text nếu chưa có summary từ Vision API
                if USE_AI and ai_service and text and len(text.strip()) > 50 and not summary:
                    try:
                        print(f"[AI Summary] Bắt đầu tóm tắt tài liệu: {page_count} trang, text length: {len(text)} ký tự")
                        # Truyền page_count để AI service biết đây là tài liệu dài
                        summary, keywords = ai_service.summarize_content(text, page_count=page_count)
                        print(f"[AI Summary] Tóm tắt thành công: summary length={len(summary) if summary else 0}, keywords count={len(keywords) if keywords else 0}")
                    except Exception as e:
                        print(f"[AI Summary] Lỗi khi tóm tắt: {e}")
                        import traceback
                        traceback.print_exc()
                        summary, keywords = None, None

                # Nếu không có summary, tạo tóm tắt cơ bản dựa trên title và metadata
                if not summary:
                    if not text or len(text.strip()) < 50:
                        # PDF scan hoặc không có text: tạo tóm tắt dựa trên title
                        if page_count > 100:
                            summary = f"Tài liệu {title} ({page_count} trang). Đây là một tài liệu dài về chủ đề được đề cập trong tiêu đề. Tài liệu có thể chứa nội dung quan trọng về {title.lower()}."
                        else:
                            summary = f"Tài liệu {title} ({page_count} trang). Tài liệu về chủ đề được đề cập trong tiêu đề."
                        print(f"[AI Summary] Tạo tóm tắt cơ bản từ title: {summary[:100]}...")
                    else:
                        # Có text nhưng ngắn: sử dụng text làm summary
                        summary = (text[:1200] + "…") if text and len(text) > 1200 else (text or f"Tài liệu: {title}.")
                if not keywords:
                    seed = (text or f"{title}").lower()
                    keywords = _naive_keywords(seed, 12)

                final_img = image_url
                if not final_img and pdf_bytes:
                    buf = _generate_thumb_from_pdf_bytes(pdf_bytes, 1.0)
                    if buf:
                        img_key = f"images/auto_thumb_{uuid.uuid4()}.jpg"
                        final_img = aws_service.upload_file(buf, img_key, "image/jpeg")

                page_count = _get_pdf_page_count(pdf_bytes) if pdf_bytes else 0

                update_fields = {"summary": summary, "keywords": keywords, "image_url": final_img}
                if page_count:
                    update_fields["pages"] = page_count

                mongo_collections.documents.update_one(
                    {"_id": doc_id},
                    {"$set": update_fields}
                )
            except Exception as e:
                print("[bg_enrich] lỗi:", e)

        concurrent.futures.ThreadPoolExecutor(max_workers=1).submit(_bg_enrich)

         # cộng điểm + ghi transaction (không chặn)
        try:
            mongo_collections.users.update_one({"_id": current_user_oid}, {"$inc": {"points": 1}})
            mongo_collections.point_txns.insert_one({
                "userId": current_user_oid,
                "type": "upload",
                "points": +1,
                "meta": {"documentId": str(doc_id)},
                "createdAt": datetime.utcnow()
            })
        except Exception as e:
            print("[register_document] lỗi cộng điểm:", e)

        return jsonify({"document_id": str(doc_id)}), 200
    except Exception as e:
        print(f"[ERROR] register_document: {e}")
        return jsonify({"error": f"Lỗi server nội bộ: {e}"}), 500


# ===================== UPLOAD (giữ route cũ, đã tối ưu A-tweaks) =====================

def _apply_rate_limit_if_available(route_func):
    """Apply rate limiting nếu flask-limiter có sẵn"""
    try:
        from flask import current_app
        limiter = current_app.config.get('LIMITER')
        if limiter:
            return limiter.limit("10 per minute")(route_func)
    except Exception:
        pass
    return route_func

@documents_bp.route("/upload", methods=["POST"])
@_apply_rate_limit_if_available
def upload_document():
    """
    Upload PDF/Word (+ optional image) – GIỮ CHO TƯƠNG THÍCH
    - Convert Word -> PDF (nếu không bật SKIP_WORD_CONVERSION)
    - Trích text (PyMuPDF/OCR) -> Tóm tắt AI + keywords (chỉ OCR khi cần)
    - Tạo thumbnail tự sinh (nếu có PDF) hoặc placeholder
    - Upload S3 & Lưu Mongo (cộng điểm user)
    """
    try:
        if "file" not in request.files:
            return jsonify({"error": "Thiếu file tài liệu (field 'file')."}), 400

        up_file = request.files["file"]
        image = request.files.get("image")
        title = (request.form.get("title") or "").strip()
        school_id = (request.form.get("schoolId") or "").strip()
        category_id = (request.form.get("categoryId") or "").strip()

        # Validate title với validation utilities
        try:
            from app.utils.validation import validate_title
            title = validate_title(title)
        except ValueError as e:
            return jsonify({"error": str(e)}), 400

        school_id, category_id = _ensure_lookup_ids(school_id, category_id)
        if up_file.filename == "" or not _allowed_ext(up_file.filename, ALLOWED_DOC_EXT):
            return jsonify({"error": "Chỉ chấp nhận PDF hoặc Word (.doc/.docx)."}), 400
        if image and image.filename and not _allowed_ext(image.filename, ALLOWED_IMG_EXT):
            return jsonify({"error": "Ảnh chỉ chấp nhận PNG/JPG/JPEG/WEBP."}), 400

        # Xác thực JWT
        try:
            current_user_oid, uploader_name = _get_current_user_strict()
        except ExpiredSignatureError:
            return jsonify({"error": "JWT hết hạn. Vui lòng đăng nhập lại."}), 401
        except InvalidTokenError as e:
            return jsonify({"error": f"Token không hợp lệ: {e}"}), 401

        # Đọc bytes & chuẩn bị
        raw_bytes = up_file.read()
        ext = os.path.splitext(secure_filename(up_file.filename))[1].lower().lstrip(".") or "pdf"

        # Convert Word -> PDF (tùy ENV)
        pdf_bytes = None
        conversion_warning = None
        if ext in ("docx", "doc") and not SKIP_WORD_CONVERSION:
            pdf_bytes = _convert_word_to_pdf_bytes(raw_bytes, ext)
            if not pdf_bytes:
                conversion_warning = "Không thể chuyển Word sang PDF: đã upload file gốc và sinh thumbnail placeholder."
        elif ext == "pdf":
            pdf_bytes = raw_bytes
        else:
            pdf_bytes = None

        # Trích text – chỉ OCR khi cần
        text = ""
        page_count = 0
        if pdf_bytes:
            # Sử dụng hàm smart để trích text từ nhiều phần cho tài liệu dài
            page_count = _get_pdf_page_count(pdf_bytes)
            print(f"[Upload] Tài liệu có {page_count} trang")
            if page_count > 100:
                # Tài liệu quá dài (>100 trang): chỉ lấy 50 trang đầu để tóm tắt
                text = _extract_text_from_pdf_bytes_smart(pdf_bytes, max_pages=50)
            else:
                # Tài liệu ngắn: lấy tất cả hoặc 6 trang đầu
                text = _extract_text_from_pdf_bytes(pdf_bytes, max_pages=min(6, page_count))
            if not text or len(text.strip()) < 100:
                print(f"[Upload] Text quá ngắn ({len(text) if text else 0} ký tự), thử OCR...")
                # Với tài liệu dài, OCR nhiều trang hơn để có đủ nội dung
                ocr_pages = 10 if page_count > 100 else 5
                ocr_text = _ocr_text_from_pdf_bytes(pdf_bytes, pages_max=ocr_pages, scale=2.0)
                if ocr_text and len(ocr_text.strip()) > len(text.strip() if text else ""):
                    text = ocr_text
                    print(f"[Upload] OCR thành công, sử dụng text từ OCR: {len(text)} ký tự")
                elif not text:
                    text = ocr_text  # Sử dụng OCR text ngay cả khi ngắn nếu không có text nào
                
                # Nếu vẫn không có text (PDF scan, tesseract không có), thử Gemini Vision API
                if (not text or len(text.strip()) < 50) and USE_AI and ai_service:
                    print(f"[Upload] Thử sử dụng Gemini Vision API để OCR và tóm tắt trực tiếp...")
                    try:
                        vision_summary, vision_keywords = ai_service.extract_and_summarize_from_pdf_images(
                            pdf_bytes, max_pages=10, page_count=page_count
                        )
                        if vision_summary:
                            # Sử dụng summary từ Gemini Vision, không cần text nữa
                            summary = vision_summary
                            keywords = vision_keywords
                            print(f"[Upload] Gemini Vision thành công: summary={len(vision_summary)} ký tự")
                            # Đặt text rỗng vì đã có summary từ Vision API
                            text = ""  # Không cần text nữa vì đã có summary
                    except Exception as e:
                        print(f"[Upload] Lỗi khi dùng Gemini Vision: {e}")
                        import traceback
                        traceback.print_exc()
        elif ext == "docx":
            text = _extract_text_from_docx_bytes(raw_bytes)
            page_count = 0  # Không có page count cho docx

        print(f"[Upload] Text đã trích xuất: {len(text) if text else 0} ký tự")

        # Tóm tắt & keywords - không cần giới hạn snippet nữa vì AI service đã xử lý
        # Lưu ý: summary có thể đã được tạo từ Gemini Vision API ở trên
        if 'summary' not in locals():
            summary, keywords = None, None
        
        # Chỉ tóm tắt bằng text nếu chưa có summary từ Vision API
        if USE_AI and ai_service is not None and text and len(text.strip()) > 50 and not summary:
            try:
                print(f"[AI Summary] Bắt đầu tóm tắt tài liệu: {page_count} trang, text length: {len(text)} ký tự")
                # Truyền page_count để AI service biết đây là tài liệu dài
                summary, keywords = ai_service.summarize_content(text, page_count=page_count)
                print(f"[AI Summary] Tóm tắt thành công: summary length={len(summary) if summary else 0}, keywords count={len(keywords) if keywords else 0}")
                if summary:
                    print(f"[AI Summary] Preview: {summary[:200]}...")
            except Exception as e:
                print(f"[AI Summary] Lỗi khi tóm tắt: {e}")
                import traceback
                traceback.print_exc()
                pass
        
        # Nếu không có summary, tạo tóm tắt cơ bản dựa trên title và metadata
        if not summary:
            if not text or len(text.strip()) < 50:
                # PDF scan hoặc không có text: tạo tóm tắt dựa trên title
                if page_count > 100:
                    summary = f"Tài liệu {title} ({page_count} trang). Đây là một tài liệu dài về chủ đề được đề cập trong tiêu đề. Tài liệu có thể chứa nội dung quan trọng về {title.lower()}."
                else:
                    summary = f"Tài liệu {title} ({page_count} trang). Tài liệu về chủ đề được đề cập trong tiêu đề."
                print(f"[AI Summary] Tạo tóm tắt cơ bản từ title: {summary[:100]}...")
            else:
                # Có text nhưng ngắn: sử dụng text làm summary
                summary = (text[:1200] + "…") if text and len(text) > 1200 else (text or f"Tài liệu: {title}.")
        if not keywords:
            seed = (text or f"{title} {up_file.filename}").lower()
            keywords = _naive_keywords(seed, 12)

        image_url = None
        # ==== Song song: upload + render thumb ====
        with concurrent.futures.ThreadPoolExecutor(max_workers=4) as ex:
            thumb_future = ex.submit(_generate_thumb_from_pdf_bytes, pdf_bytes, 1.0) if pdf_bytes else None

            if pdf_bytes:
                pdf_key = f"documents/{uuid.uuid4()}.pdf"
                upload_future = ex.submit(aws_service.upload_file, BytesIO(pdf_bytes), pdf_key, "application/pdf")
            else:
                obj_ext = "." + ext
                obj_ct = up_file.mimetype or (
                    "application/msword" if ext == "doc" else
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                doc_key = f"documents/{uuid.uuid4()}{obj_ext}"
                upload_future = ex.submit(aws_service.upload_file, BytesIO(raw_bytes), doc_key, obj_ct)

            s3_url = upload_future.result()
            if not s3_url:
                return jsonify({"error": "Upload lên S3 thất bại."}), 500

            if image and image.filename:
                try:
                    img_ext = os.path.splitext(secure_filename(image.filename))[1].lower() or ".jpg"
                    img_key = f"images/{uuid.uuid4()}{img_ext}"
                    img_ct = image.mimetype or "image/jpeg"
                    image_url = aws_service.upload_file(image.stream, img_key, img_ct)
                except Exception:
                    image_url = None

            if not image_url and thumb_future:
                thumb_buf = thumb_future.result()
                if thumb_buf:
                    thumb_key = f"images/auto_thumb_{uuid.uuid4()}.jpg"
                    image_url = aws_service.upload_file(thumb_buf, thumb_key, "image/jpeg")

        page_count = _get_pdf_page_count(pdf_bytes) if pdf_bytes else 0

        # Lưu Mongo
        doc = Document(
            title=title,
            s3_url=s3_url,
            user_id=str(current_user_oid),
            summary=summary,
            keywords=keywords or [],
            school_id=str(school_id),
            category_id=str(category_id),
            image_url=image_url,
            pages=page_count if page_count else None,
        )
        doc_dict = doc.to_mongo_doc()
        doc_dict["uploaderName"] = uploader_name
        result = mongo_collections.documents.insert_one(doc_dict)

        # Cộng điểm + ghi transaction
        try:
            mongo_collections.users.update_one({"_id": current_user_oid}, {"$inc": {"points": 1}})
            mongo_collections.point_txns.insert_one({
                "userId": current_user_oid,
                "type": "upload",
                "points": +1,
                "meta": {"documentId": str(result.inserted_id)},
                "createdAt": datetime.utcnow()
            })
        except Exception as e:
            print("[upload_document] lỗi cộng điểm:", e)


        payload = {
            "message": "Upload thành công",
            "document_id": str(result.inserted_id),
            "s3_url": s3_url,
            "image_url": image_url,
            "summary": summary,
            "keywords": keywords,
        }
        if conversion_warning:
            payload["warning"] = conversion_warning
        return jsonify(payload), 200

    except Exception as e:
        print(f"[ERROR] upload_document: {e}")
        return jsonify({"error": f"Lỗi server nội bộ: {e}"}), 500


# ===================== LIST =====================

@documents_bp.route("", methods=["GET"])
@documents_bp.route("/", methods=["GET"])
def get_documents():
    """
    Lấy danh sách documents với search + filters + pagination.
    Hỗ trợ: search, schoolId, categoryId, fileType, length (short/medium/long),
            uploadDate (today|yesterday|last7days|last30days|month:YYYY:MM|year:YYYY|day:YYYY:MM:DD|week:YYYY:WW)
            page (default: 1), limit (default: 12, max: 100)
    """
    try:
        search = (request.args.get("search") or "").strip()
        school_id = (request.args.get("schoolId") or "").strip()
        category_id = (request.args.get("categoryId") or "").strip()
        file_type = (request.args.get("fileType") or "").strip().lower()
        length = (request.args.get("length") or "").strip().lower()
        upload_date = (request.args.get("uploadDate") or "").strip()
        
        # Pagination parameters
        try:
            page = int(request.args.get("page", 1))
            if page < 1:
                page = 1
        except (ValueError, TypeError):
            page = 1
        
        try:
            limit = int(request.args.get("limit", 12))
            if limit < 1:
                limit = 12
            if limit > 100:
                limit = 100
        except (ValueError, TypeError):
            limit = 12
        
        skip = (page - 1) * limit

        query_or = []
        if search:
            query_or = [
                {"title": {"$regex": search, "$options": "i"}},
                {"summary": {"$regex": search, "$options": "i"}},
                {"keywords": {"$regex": search, "$options": "i"}},
            ]

        ands = []

        # school/category: hỗ trợ cả ObjectId lẫn string (tương thích dữ liệu cũ)
        def _or_id(field1, field2, val):
            ors = []
            try:
                ors += [{field1: ObjectId(val)}, {field2: ObjectId(val)}]
            except Exception:
                pass
            ors += [{field1: val}, {field2: val}]
            return {"$or": ors}

        if school_id:
            ands.append(_or_id("schoolId", "school_id", school_id))
        if category_id:
            ands.append(_or_id("categoryId", "category_id", category_id))

        # file type theo đuôi URL
        if file_type == "pdf":
            ands.append({"s3_url": {"$regex": r"\.pdf$", "$options": "i"}})
        elif file_type in {"doc", "docx", "word"}:
            ands.append({"s3_url": {"$regex": r"\.(docx|doc)$", "$options": "i"}})

        # length (pages)
        if length == "short":
            ands.append({"pages": {"$lt": 10}})
        elif length == "medium":
            ands.append({"pages": {"$gte": 10, "$lte": 50}})
        elif length == "long":
            ands.append({"pages": {"$gt": 50}})

        # upload date
        if upload_date:
            now = datetime.utcnow()
            today_start = datetime(now.year, now.month, now.day)
            date_filter = None
            try:
                if upload_date == "today":
                    date_filter = {"$gte": today_start}
                elif upload_date == "yesterday":
                    ys = today_start - timedelta(days=1)
                    date_filter = {"$gte": ys, "$lt": today_start}
                elif upload_date == "last7days":
                    date_filter = {"$gte": now - timedelta(days=7)}
                elif upload_date == "last30days":
                    date_filter = {"$gte": now - timedelta(days=30)}
                elif upload_date.startswith("month:"):
                    _, y, m = upload_date.split(":")
                    y, m = int(y), int(m)
                    m0 = datetime(y, m, 1)
                    m1 = datetime(y+1, 1, 1) if m == 12 else datetime(y, m+1, 1)
                    date_filter = {"$gte": m0, "$lt": m1}
                elif upload_date.startswith("year:"):
                    _, y = upload_date.split(":")
                    y = int(y)
                    y0, y1 = datetime(y, 1, 1), datetime(y+1, 1, 1)
                    date_filter = {"$gte": y0, "$lt": y1}
                elif upload_date.startswith("day:"):
                    _, y, m, d = upload_date.split(":")
                    d0 = datetime(int(y), int(m), int(d))
                    d1 = d0 + timedelta(days=1)
                    date_filter = {"$gte": d0, "$lt": d1}
                elif upload_date.startswith("week:"):
                    _, y, w = upload_date.split(":")
                    y, w = int(y), int(w)
                    jan4 = date(y, 1, 4)
                    monday = jan4 - timedelta(days=jan4.weekday())
                    start = monday + timedelta(weeks=w-1)
                    week_start = datetime.combine(start, datetime.min.time())
                    week_end = week_start + timedelta(days=7)
                    date_filter = {"$gte": week_start, "$lt": week_end}
            except Exception:
                date_filter = None

            if date_filter:
                ands.append({"$or": [{"createdAt": date_filter}, {"created_at": date_filter}]})

        # Build mongo query
        mongo_query = {}
        if query_or:
            mongo_query["$or"] = query_or
        if ands:
            if "$or" in mongo_query:
                mongo_query = {"$and": [{"$or": mongo_query["$or"]}] + ands}
            else:
                mongo_query = {"$and": ands}

        # Query & sort với pagination
        try:
            cursor = mongo_collections.documents.find(mongo_query).sort("createdAt", -1)
        except Exception:
            try:
                cursor = mongo_collections.documents.find(mongo_query).sort("created_at", -1)
            except Exception:
                cursor = mongo_collections.documents.find(mongo_query)

        # Lấy tổng số documents (luôn tính để frontend có thể hiển thị đúng)
        total_count = None
        try:
            # Sử dụng count_documents thay vì cursor.count() (deprecated)
            total_count = mongo_collections.documents.count_documents(mongo_query)
        except Exception:
            pass
        
        # Apply pagination
        docs = list(cursor.skip(skip).limit(limit))

        # Chuẩn bị map thống kê phản ứng/bình luận
        doc_ids = [doc.get("_id") for doc in docs if doc.get("_id")]
        reaction_map = {}
        comment_map = {}
        if doc_ids:
            try:
                reaction_pipeline = [
                    {"$match": {"documentId": {"$in": doc_ids}}},
                    {
                        "$group": {
                            "_id": "$documentId",
                            "likes": {
                                "$sum": {
                                    "$cond": [
                                        {"$eq": ["$reaction", "like"]},
                                        1,
                                        0,
                                    ]
                                }
                            },
                            "dislikes": {
                                "$sum": {
                                    "$cond": [
                                        {"$eq": ["$reaction", "dislike"]},
                                        1,
                                        0,
                                    ]
                                }
                            },
                        }
                    },
                ]
                for row in mongo_collections.document_reactions.aggregate(reaction_pipeline):
                    reaction_map[str(row.get("_id"))] = {
                        "likes": row.get("likes", 0),
                        "dislikes": row.get("dislikes", 0),
                    }
            except Exception as agg_err:
                print("[get_documents] aggregate reactions error:", agg_err)

            try:
                comment_pipeline = [
                    {"$match": {"documentId": {"$in": doc_ids}}},
                    {"$group": {"_id": "$documentId", "count": {"$sum": 1}}},
                ]
                for row in mongo_collections.document_comments.aggregate(comment_pipeline):
                    comment_map[str(row.get("_id"))] = row.get("count", 0)
            except Exception as agg_err:
                print("[get_documents] aggregate comments error:", agg_err)

        # Tối ưu: Load tất cả schools/categories/users một lần thay vì N queries
        school_ids = set()
        category_ids = set()
        user_ids = set()
        pages_to_update = []  # Collect updates để thực hiện sau
        
        for doc in docs:
            sid = doc.get("schoolId") or doc.get("school_id")
            if sid:
                try:
                    school_ids.add(sid if isinstance(sid, ObjectId) else ObjectId(str(sid)))
                except Exception:
                    pass
            
            cid = doc.get("categoryId") or doc.get("category_id")
            if cid:
                try:
                    category_ids.add(cid if isinstance(cid, ObjectId) else ObjectId(str(cid)))
                except Exception:
                    pass
            
            uid = doc.get("userId") or doc.get("user_id")
            if uid:
                try:
                    user_ids.add(uid if isinstance(uid, ObjectId) else ObjectId(str(uid)))
                except Exception:
                    pass

        # Load tất cả schools/categories/users trong 1 query mỗi loại
        school_map = {}
        if school_ids:
            for s in mongo_collections.schools.find({"_id": {"$in": list(school_ids)}}):
                school_map[str(s["_id"])] = {"_id": str(s["_id"]), "name": s.get("name", "")}
        
        category_map = {}
        if category_ids:
            for c in mongo_collections.categories.find({"_id": {"$in": list(category_ids)}}):
                category_map[str(c["_id"])] = {"_id": str(c["_id"]), "name": c.get("name", "")}
        
        user_map = {}
        if user_ids:
            for u in mongo_collections.users.find({"_id": {"$in": list(user_ids)}}):
                user_map[str(u["_id"])] = {
                    "_id": str(u["_id"]),
                    "username": u.get("username", ""),
                    "name": u.get("fullName") or u.get("name", "")
                }

        # Map kết quả
        result = []
        for doc in docs:
            doc_id_obj = doc.get("_id")
            doc_id_str = str(doc_id_obj) if doc_id_obj else None

            # chấp nhận nhiều key pages khác nhau
            raw_pages_candidates = [
                doc.get("pages"),
                doc.get("pageCount"),
                doc.get("page_count"),
            ]
            metadata = doc.get("metadata")
            if isinstance(metadata, dict):
                raw_pages_candidates.append(metadata.get("pages"))

            pages_val = None
            for candidate in raw_pages_candidates:
                if candidate is None or candidate == "":
                    continue
                try:
                    pages_val = int(candidate)
                    break
                except (TypeError, ValueError):
                    continue
            if pages_val is None:
                pages_val = 0

            # Collect updates để thực hiện sau (không block response)
            if doc_id_obj and ("pages" not in doc or doc.get("pages") in (None, "")):
                inferred_pages = _infer_document_page_count(doc)
                if inferred_pages:
                    pages_val = inferred_pages
                    pages_to_update.append((doc_id_obj, inferred_pages))
                else:
                    pages_to_update.append((doc_id_obj, 0))

            reactions = reaction_map.get(doc_id_str, {})
            likes = reactions.get("likes", 0)
            dislikes = reactions.get("dislikes", 0)
            comment_count = comment_map.get(doc_id_str, 0)

            out = {
                "_id": doc_id_str,
                "id": doc_id_str,
                "title": doc.get("title", ""),
                "summary": doc.get("summary", ""),
                "s3_url": doc.get("s3_url", ""),
                "image_url": doc.get("image_url"),
                "keywords": doc.get("keywords", []),
                "created_at": (doc.get("created_at") or doc.get("createdAt")).isoformat()
                               if (doc.get("created_at") or doc.get("createdAt")) else None,
                "createdAt": (doc.get("created_at") or doc.get("createdAt")).isoformat()
                               if (doc.get("created_at") or doc.get("createdAt")) else None,
                "upload_date": (doc.get("created_at") or doc.get("createdAt")).isoformat()
                               if (doc.get("created_at") or doc.get("createdAt")) else None,
                "views": doc.get("views", 0),
                "pages": pages_val,
                "likes": likes,
                "dislikes": dislikes,
                "commentCount": comment_count,
                "totalReviews": likes + dislikes + comment_count,
            }

            # Join school/category/user từ maps đã load
            sid = doc.get("schoolId") or doc.get("school_id")
            if sid:
                try:
                    sid_str = str(sid) if isinstance(sid, ObjectId) else str(ObjectId(str(sid)))
                    s = school_map.get(sid_str)
                    if s:
                        out["school"] = s
                        out["school_name"] = s.get("name", "")
                except Exception:
                    pass

            cid = doc.get("categoryId") or doc.get("category_id")
            if cid:
                try:
                    cid_str = str(cid) if isinstance(cid, ObjectId) else str(ObjectId(str(cid)))
                    c = category_map.get(cid_str)
                    if c:
                        out["category"] = c
                        out["category_name"] = c.get("name", "")
                except Exception:
                    pass

            uid = doc.get("userId") or doc.get("user_id")
            if uid:
                try:
                    uid_str = str(uid) if isinstance(uid, ObjectId) else str(ObjectId(str(uid)))
                    u = user_map.get(uid_str)
                    if u:
                        out["user"] = u
                        out["uploader"] = u.get("name") or u.get("username", "")
                except Exception:
                    pass

            result.append(out)
        
        # Thực hiện updates sau khi đã trả response (async, không block)
        if pages_to_update:
            try:
                from threading import Thread
                def update_pages():
                    for doc_id, pages_val in pages_to_update:
                        try:
                            mongo_collections.documents.update_one(
                                {"_id": doc_id},
                                {"$set": {"pages": pages_val}},
                            )
                        except Exception:
                            pass
                Thread(target=update_pages, daemon=True).start()
            except Exception:
                pass

        # Trả về với pagination metadata nếu có total_count
        response = result
        if total_count is not None:
            response = {
                "documents": result,
                "total": total_count,
                "page": page,
                "limit": limit,
                "totalPages": (total_count + limit - 1) // limit if limit > 0 else 0
            }
        
        return jsonify(response), 200

    except Exception as e:
        print(f"[ERROR] get_documents: {e}")
        import traceback; traceback.print_exc()
        return jsonify({"error": f"Lỗi server nội bộ: {e}"}), 500


# ===================== DETAIL =====================

@documents_bp.route("/<doc_id>", methods=["GET"])
def get_document_detail(doc_id):
    """Lấy chi tiết 1 tài liệu theo id."""
    try:
        _id = ObjectId(doc_id)
    except Exception:
        return jsonify({"error": "document id không hợp lệ"}), 400

    proj = {
        "title": 1, "summary": 1, "keywords": 1, "image_url": 1, "s3_url": 1,
        "schoolId": 1, "categoryId": 1, "userId": 1, "uploaderName": 1,
        "createdAt": 1, "created_at": 1
    }
    d = mongo_collections.documents.find_one({"_id": _id}, proj)
    if not d:
        return jsonify({"error": "Không tìm thấy tài liệu"}), 404

    # School
    school_name = None
    sid = d.get("schoolId")
    if sid:
        try:
            s = mongo_collections.schools.find_one({"_id": sid if isinstance(sid, ObjectId) else ObjectId(str(sid))}, {"name": 1})
            if s:
                school_name = s.get("name")
        except Exception:
            pass

    # Category
    category_name = None
    cid = d.get("categoryId")
    if cid:
        try:
            c = mongo_collections.categories.find_one({"_id": cid if isinstance(cid, ObjectId) else ObjectId(str(cid))}, {"name": 1})
            if c:
                category_name = c.get("name")
        except Exception:
            pass

    # Uploader
    uploader = d.get("uploaderName")
    uploader_id_val = d.get("userId") or d.get("user_id")
    if not uploader and uploader_id_val:
        try:
            u = mongo_collections.users.find_one(
                {"_id": uploader_id_val if isinstance(uploader_id_val, ObjectId) else ObjectId(str(uploader_id_val))},
                {"fullName": 1, "username": 1, "name": 1, "email": 1}
            )
            if u:
                uploader = u.get("fullName") or u.get("username") or u.get("name") or u.get("email")
        except Exception:
            pass

    created = d.get("createdAt") or d.get("created_at")
    page_count = (
        d.get("pages")
        or d.get("pageCount")
        or d.get("page_count")
        or (d.get("metadata", {}).get("pages") if isinstance(d.get("metadata"), dict) else None)
    )
    try:
        page_count = int(page_count) if page_count is not None else None
    except (TypeError, ValueError):
        page_count = None
    return jsonify({
        "_id": str(d["_id"]),
        "title": d.get("title"),
        "summary": d.get("summary"),
        "keywords": d.get("keywords") or [],
        "image_url": d.get("image_url"),
        "s3_url": d.get("s3_url"),
        "schoolName": school_name,
        "categoryName": category_name,
        "uploaderName": uploader,
        "uploaderId": str(uploader_id_val) if uploader_id_val else None,
        "createdAt": created.isoformat() if created else None,
        "pages": page_count,
    }), 200


@documents_bp.route("/<string:doc_id>/reactions", methods=["GET"])
def get_document_reactions(doc_id):
    try:
        doc_oid = ObjectId(doc_id)
    except Exception:
        return jsonify({"error": "document id không hợp lệ"}), 400

    like_count = mongo_collections.document_reactions.count_documents({"documentId": doc_oid, "reaction": "like"})
    dislike_count = mongo_collections.document_reactions.count_documents({"documentId": doc_oid, "reaction": "dislike"})

    user_id, _ = _get_current_user_optional()
    my_reaction = None
    if user_id:
        doc = mongo_collections.document_reactions.find_one({"documentId": doc_oid, "userId": user_id}, {"reaction": 1})
        if doc:
            my_reaction = doc.get("reaction")

    return jsonify({
        "likes": like_count,
        "dislikes": dislike_count,
        "myReaction": my_reaction,
    }), 200


@documents_bp.route("/<string:doc_id>/reactions", methods=["POST"])
def set_document_reaction(doc_id):
    try:
        doc_oid = ObjectId(doc_id)
    except Exception:
        return jsonify({"error": "document id không hợp lệ"}), 400

    try:
        user_id, _ = _get_current_user_strict()
    except ExpiredSignatureError:
        return jsonify({"error": "JWT hết hạn. Vui lòng đăng nhập lại."}), 401
    except InvalidTokenError as e:
        return jsonify({"error": f"Token không hợp lệ: {e}"}), 401

    data = request.get_json() or {}
    action = (data.get("action") or "").lower()
    if action not in {"like", "dislike", "none"}:
        return jsonify({"error": "action phải là like/dislike/none"}), 400

    if action == "none":
        mongo_collections.document_reactions.delete_one({"documentId": doc_oid, "userId": user_id})
    else:
        mongo_collections.document_reactions.update_one(
            {"documentId": doc_oid, "userId": user_id},
            {"$set": {"reaction": action, "updatedAt": datetime.utcnow()}},
            upsert=True,
        )

    like_count = mongo_collections.document_reactions.count_documents({"documentId": doc_oid, "reaction": "like"})
    dislike_count = mongo_collections.document_reactions.count_documents({"documentId": doc_oid, "reaction": "dislike"})

    return jsonify({
        "likes": like_count,
        "dislikes": dislike_count,
        "myReaction": None if action == "none" else action,
    }), 200


@documents_bp.route("/<string:doc_id>/comments", methods=["GET"])
def list_document_comments(doc_id):
    try:
        doc_oid = ObjectId(doc_id)
    except Exception:
        return jsonify({"error": "document id không hợp lệ"}), 400

    cursor = mongo_collections.document_comments.find({"documentId": doc_oid}).sort("createdAt", 1).limit(200)
    comments = list(cursor)

    user_ids = {c.get("userId") for c in comments if c.get("userId")}
    users_map = {}
    if user_ids:
        cursor_users = mongo_collections.users.find(
            {"_id": {"$in": list(user_ids)}},
            {"username": 1, "fullName": 1, "avatarUrl": 1}
        )
        for u in cursor_users:
            users_map[str(u["_id"])] = {
                "id": str(u["_id"]),
                "username": u.get("username"),
                "fullName": u.get("fullName"),
                "avatarUrl": u.get("avatarUrl"),
            }

    payload = []
    for c in comments:
        payload.append({
            "id": str(c.get("_id")),
            "content": c.get("content", ""),
            "createdAt": (c.get("createdAt") or datetime.utcnow()).isoformat() + "Z",
            "user": users_map.get(str(c.get("userId")), {
                "id": str(c.get("userId")) if c.get("userId") else None,
                "username": None,
                "fullName": None,
                "avatarUrl": None,
            })
        })

    return jsonify(payload), 200


@documents_bp.route("/<string:doc_id>/comments", methods=["POST"])
def create_document_comment(doc_id):
    try:
        doc_oid = ObjectId(doc_id)
    except Exception:
        return jsonify({"error": "document id không hợp lệ"}), 400

    try:
        user_id, user_doc = _get_current_user_strict(return_doc=True)
    except ExpiredSignatureError:
        return jsonify({"error": "JWT hết hạn. Vui lòng đăng nhập lại."}), 401
    except InvalidTokenError as e:
        return jsonify({"error": f"Token không hợp lệ: {e}"}), 401

    data = request.get_json() or {}
    content = (data.get("content") or "").strip()
    if not content:
        return jsonify({"error": "Nội dung bình luận không được trống"}), 400
    if len(content) > 1000:
        return jsonify({"error": "Nội dung bình luận quá dài (tối đa 1000 ký tự)"}), 400

    comment_doc = {
        "documentId": doc_oid,
        "userId": user_id,
        "content": content,
        "createdAt": datetime.utcnow(),
    }
    result = mongo_collections.document_comments.insert_one(comment_doc)

    payload = {
        "id": str(result.inserted_id),
        "content": content,
        "createdAt": comment_doc["createdAt"].isoformat() + "Z",
        "user": {
            "id": str(user_id),
            "username": user_doc.get("username"),
            "fullName": user_doc.get("fullName"),
            "avatarUrl": user_doc.get("avatarUrl"),
        },
    }

    return jsonify(payload), 201


# ===================== RAW (Proxy PDF cho pdf.js) =====================

def _parse_s3_url(s3_url: str):
    """
    Trả về (bucket, key) từ S3 URL:
      - https://mybucket.s3.ap-southeast-1.amazonaws.com/path/to/file.pdf
      - https://s3.ap-southeast-1.amazonaws.com/mybucket/path/to/file.pdf
    """
    try:
        u = urlparse(s3_url)
        host = u.netloc
        path = u.path.lstrip("/")
        if host.startswith("s3.") or host.startswith("s3-"):
            parts = path.split("/", 1)
            if len(parts) == 2:
                return parts[0], parts[1]
        if ".s3." in host:
            bucket = host.split(".s3.")[0]
            return bucket, path
    except Exception:
        pass
    return None, None


def _stream_boto3(bucket: str, key: str, range_header: str | None):
    """Stream object từ S3 bằng boto3, hỗ trợ Range (cần cho pdf.js)."""
    s3 = boto3.client(
        "s3",
        region_name=os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION"),
        aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
        aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY")
    )
    params = {"Bucket": bucket, "Key": key}
    if range_header:
        params["Range"] = range_header

    obj = s3.get_object(**params)
    body = obj["Body"]

    def gen():
        for chunk in body.iter_chunks(chunk_size=8192):
            if chunk:
                yield chunk

    headers = {
        "Content-Type": obj.get("ContentType", "application/pdf"),
        "Accept-Ranges": "bytes",
    }
    if "ContentLength" in obj and not range_header:
        headers["Content-Length"] = str(obj["ContentLength"])
    if "ContentRange" in obj:
        headers["Content-Range"] = obj["ContentRange"]
    status = 206 if "Content-Range" in headers else 200
    return Response(stream_with_context(gen()), status=status, headers=headers)


# @documents_bp.route("/<string:doc_id>/raw", methods=["GET"])
# def get_document_raw(doc_id):
#     """
#     Proxy PDF về FE:
#       1) Thử stream trực tiếp theo URL đã lưu (nếu public/presigned)
#       2) Nếu thất bại -> bóc bucket/key và stream qua boto3 (không cần URL public)
#     """
#     try:
#         _id = ObjectId(doc_id)
#     except Exception:
#         return jsonify({"error": "document id không hợp lệ"}), 400

#     d = mongo_collections.documents.find_one({"_id": _id}, {"s3_url": 1})
#     if not d or not d.get("s3_url"):
#         return jsonify({"error": "Không có file PDF"}), 404

#     s3_url = d["s3_url"]
#     range_h = request.headers.get("Range")

#     # 1) Thử stream theo URL đã lưu
#     try:
#         fwd = {"Range": range_h} if range_h else {}
#         r = requests.get(s3_url, headers=fwd, stream=True, timeout=15)
#         if r.status_code < 400:
#             ctype = (r.headers.get("Content-Type") or "").lower()
#             if "application/pdf" in ctype or "octet-stream" in ctype:
#                 def gen_req():
#                     for chunk in r.iter_content(chunk_size=8192):
#                         if chunk:
#                             yield chunk
#                 resp = Response(stream_with_context(gen_req()), status=r.status_code)
#                 for h in ["Content-Type", "Content-Length", "Accept-Ranges", "Content-Range"]:
#                     v = r.headers.get(h)
#                     if v:
#                         resp.headers[h] = v
#                 resp.headers["X-Content-Type-Options"] = "nosniff"
#                 return resp
#     except Exception:
#         pass

#     # 2) Fallback: stream qua boto3
#     bucket, key = _parse_s3_url(s3_url)
#     if not bucket or not key:
#         return jsonify({"error": "Không parse được S3 URL (bucket/key)."}), 500

#     try:
#         return _stream_boto3(bucket, key, range_h)
#     except Exception as e:
#         return jsonify({"error": f"Không thể đọc PDF từ S3: {str(e)}"}), 502
# # ... phía trên đã có import fitz, PIL, requests, v.v.

@documents_bp.route("/<string:doc_id>/text", methods=["GET"])
def get_document_text(doc_id):
    """
    Trả về văn bản thuần đã CHUẨN HOÁ:
      - Bỏ ngắt trang (\f)
      - Bỏ dòng số trang / 'Page x / y'
      - Ghép từ bị gạch nối ở cuối dòng
      - Gộp các dòng trong cùng đoạn thành 1 dòng (giữ khoảng trống giữa đoạn)
    Response: { text: "..." }
    """
    def _normalize_text(raw: str) -> str:
        import re
        if not raw:
            return ""

        # 1) thống nhất xuống dòng, bỏ form-feed (ngắt trang)
        s = raw.replace("\r\n", "\n").replace("\r", "\n")
        s = s.replace("\x0c", "\n").replace("\f", "\n")  # page breaks

        # 2) bỏ header/footer kiểu "Page 1 of 69", "1 / 69", hoặc dòng chỉ có số
        #   - xử lý theo dòng để an toàn
        lines = s.split("\n")
        cleaned = []
        page_re_1 = re.compile(r"^\s*page\s+\d+(\s*(/|of)\s*\d+)?\s*$", re.IGNORECASE)
        page_re_2 = re.compile(r"^\s*\d+\s*/\s*\d+\s*$")
        page_re_3 = re.compile(r"^\s*\d+\s*$")  # dòng chỉ là số
        for ln in lines:
            l = ln.strip()
            if not l:
                cleaned.append(ln)
                continue
            if page_re_1.match(l) or page_re_2.match(l) or page_re_3.match(l):
                # bỏ dòng số trang/header/footer
                continue
            cleaned.append(ln)
        s = "\n".join(cleaned)

        # 3) xoá khoảng trắng thừa đầu/cuối dòng
        s = "\n".join([re.sub(r"\s+$", "", re.sub(r"^\s+", "", ln)) for ln in s.split("\n")])

        # 4) nối từ bị gạch nối ở cuối dòng: "thuật-\n toán" -> "thuật toán"
        s = re.sub(r"(\w)-\n(\w)", r"\1\2", s)

        # 5) gộp các dòng đơn trong cùng đoạn thành 1 dòng:
        #    - 2+ newline => giữ làm ngắt đoạn
        #    - 1 newline giữa 2 chữ => thay bằng khoảng trắng
        s = re.sub(r"[ \t]*\n[ \t]*(?=\S)", " ", s)  # newline đơn -> space
        s = re.sub(r"(?:\n\s*){2,}", "\n\n", s)     # nhiều newline -> 2 newline

        # 6) bỏ khoảng trắng thừa
        s = re.sub(r"[ \t]{2,}", " ", s)
        s = s.strip()

        return s

    from bson.objectid import ObjectId
    try:
        _id = ObjectId(doc_id)
    except Exception:
        return jsonify({"error": "document id không hợp lệ"}), 400

    d = mongo_collections.documents.find_one({"_id": _id}, {"s3_url": 1, "title": 1})
    if not d or not d.get("s3_url"):
        return jsonify({"error": "Không tìm thấy file"}), 404

    s3_url = d["s3_url"]
    try:
        r = requests.get(s3_url, timeout=45)
        if r.status_code >= 400:
            return jsonify({"error": f"Không tải được file từ S3: HTTP {r.status_code}"}), 502
        file_bytes = r.content
    except Exception as e:
        return jsonify({"error": f"Lỗi tải file: {e}"}), 502

    text = ""
    low = s3_url.lower()
    try:
        if low.endswith(".pdf"):
            text = _extract_text_from_pdf_bytes(file_bytes, max_pages=999)
            if not text:
                text = _ocr_text_from_pdf_bytes(file_bytes, pages_max=5, scale=2.0)
        elif low.endswith(".docx"):
            text = _extract_text_from_docx_bytes(file_bytes)
        elif low.endswith(".doc"):
            pdf_bytes = _convert_word_to_pdf_bytes(file_bytes, "doc")
            if pdf_bytes:
                text = _extract_text_from_pdf_bytes(pdf_bytes, max_pages=999)
    except Exception:
        text = ""

    if not text:
        title = d.get("title") or "Document"
        text = f"{title}\n\n(Chưa thể trích xuất nội dung hoặc file là ảnh scan.)"

    return jsonify({"text": _normalize_text(text)})

# app/controllers/documents.py  (chỉ hiển thị phần route /raw đã sửa)

@documents_bp.route("/<string:doc_id>/raw", methods=["GET"])
def get_document_raw(doc_id):
    """
    Proxy file về FE:
      - ?download=1 => ép tải xuống (Content-Disposition: attachment; filename="...")
      - không download => xem trước (iframe)
      - Hỗ trợ Range cho PDF
    """
    try:
        _id = ObjectId(doc_id)
    except Exception:
        return jsonify({"error": "document id không hợp lệ"}), 400

    d = mongo_collections.documents.find_one({"_id": _id}, {"s3_url": 1, "title": 1})
    if not d or not d.get("s3_url"):
        return jsonify({"error": "Không có file"}), 404

    s3_url = d["s3_url"]
    title = (d.get("title") or "document").strip()
    dl = (request.args.get("download") == "1")
    custom_name = (request.args.get("filename") or title).strip()

    range_h = request.headers.get("Range")
    safe_name = re.sub(r'[\\/:*?"<>|]+', "_", custom_name)

    # 1) Thử stream trực tiếp theo URL đã lưu
    try:
        fwd = {"Range": range_h} if range_h else {}
        r = requests.get(s3_url, headers=fwd, stream=True, timeout=20)
        if r.status_code < 400:
            ctype = r.headers.get("Content-Type") or "application/octet-stream"

            def gen_req():
                for chunk in r.iter_content(chunk_size=8192):
                    if chunk:
                        yield chunk

            resp = Response(stream_with_context(gen_req()), status=r.status_code)
            # headers cần thiết cho xem trước / range
            for h in ["Content-Type", "Content-Length", "Accept-Ranges", "Content-Range"]:
                v = r.headers.get(h)
                if v:
                    resp.headers[h] = v
            resp.headers["X-Content-Type-Options"] = "nosniff"

            # ⬇️ nếu yêu cầu download, ép Content-Disposition
            if dl:
                # đoán đuôi từ URL
                ext = os.path.splitext(urlparse(s3_url).path)[1] or ".bin"
                resp.headers["Content-Disposition"] = f'attachment; filename="{safe_name}{ext}"'
            return resp
    except Exception:
        pass

    # 2) Fallback: đọc trực tiếp từ S3 bằng boto3
    bucket, key = _parse_s3_url(s3_url)
    if not bucket or not key:
        return jsonify({"error": "Không parse được S3 URL (bucket/key)."}), 500

    try:
        s3 = boto3.client(
            "s3",
            region_name=os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION"),
            aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
            aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY")
        )
        params = {"Bucket": bucket, "Key": key}
        if range_h:
            params["Range"] = range_h
        obj = s3.get_object(**params)
        body = obj["Body"]

        def gen():
            for chunk in body.iter_chunks(chunk_size=8192):
                if chunk:
                    yield chunk

        headers = {
            "Content-Type": obj.get("ContentType", "application/octet-stream"),
            "Accept-Ranges": "bytes",
        }
        if "ContentLength" in obj and not range_h:
            headers["Content-Length"] = str(obj["ContentLength"])
        if "ContentRange" in obj:
            headers["Content-Range"] = obj["ContentRange"]

        if dl:
            ext = os.path.splitext(key)[1] or ".bin"
            headers["Content-Disposition"] = f'attachment; filename="{safe_name}{ext}"'

        status = 206 if "Content-Range" in headers else 200
        return Response(stream_with_context(gen()), status=status, headers=headers)
    except Exception as e:
        return jsonify({"error": f"Không thể đọc từ S3: {str(e)}"}), 502


# ===================== VIEW COUNT =====================

@documents_bp.route("/<string:doc_id>/view", methods=["POST"])
def increment_document_view(doc_id):
    """Tăng lượt xem của document và lưu vào lịch sử xem."""
    try:
        _id = ObjectId(doc_id)
    except Exception:
        return jsonify({"error": "document id không hợp lệ"}), 400

    try:
        # Lấy user hiện tại (nếu có token)
        current_user_oid = None
        try:
            current_user_oid, _ = _get_current_user_strict()
        except Exception:
            pass  # Cho phép xem không cần đăng nhập

        result = mongo_collections.documents.update_one(
            {"_id": _id},
            {"$inc": {"views": 1}},
            upsert=False
        )
        
        if result.matched_count == 0:
            return jsonify({"error": "Không tìm thấy tài liệu"}), 404
        
        # Lưu vào lịch sử xem nếu có user đăng nhập
        if current_user_oid:
            try:
                # Tạo hoặc cập nhật view history (chỉ lưu 1 lần gần nhất cho mỗi user-document)
                mongo_collections.view_history.update_one(
                    {"userId": current_user_oid, "documentId": _id},
                    {"$set": {"viewedAt": datetime.utcnow()}},
                    upsert=True
                )
            except Exception as e:
                print(f"[WARNING] Lỗi lưu view history: {e}")
        
        # Lấy lượt xem mới
        doc = mongo_collections.documents.find_one({"_id": _id}, {"views": 1})
        views = doc.get("views", 0) if doc else 0
        
        return jsonify({"success": True, "views": views}), 200
    except Exception as e:
        print(f"[ERROR] increment_document_view: {e}")
        return jsonify({"error": f"Lỗi server nội bộ: {e}"}), 500


# ===================== FEATURED WEEK =====================

@documents_bp.route("/featured-week", methods=["GET"])
def get_featured_documents_week():
    """Lấy các tài liệu có lượt xem nhiều nhất trong 7 ngày qua."""
    try:
        limit = int(request.args.get("limit", 3))
        if limit < 1 or limit > 10:
            limit = 3
        
        # Tính ngày bắt đầu (7 ngày trước)
        now = datetime.utcnow()
        week_start = now - timedelta(days=7)
        
        # Query: tài liệu được tạo trong 7 ngày qua, sort theo views giảm dần
        mongo_query = {
            "$or": [
                {"createdAt": {"$gte": week_start}},
                {"created_at": {"$gte": week_start}}
            ]
        }
        
        # Lấy documents và sort theo views
        cursor = mongo_collections.documents.find(mongo_query).sort("views", -1).limit(limit)
        docs = list(cursor)
        
        if not docs:
            return jsonify({"documents": []}), 200
        
        # Lấy doc_ids để aggregate reactions và comments
        doc_ids = [doc.get("_id") for doc in docs if doc.get("_id")]
        reaction_map = {}
        comment_map = {}
        
        if doc_ids:
            try:
                reaction_pipeline = [
                    {"$match": {"documentId": {"$in": doc_ids}}},
                    {
                        "$group": {
                            "_id": "$documentId",
                            "likes": {
                                "$sum": {
                                    "$cond": [
                                        {"$eq": ["$reaction", "like"]},
                                        1,
                                        0,
                                    ]
                                }
                            },
                            "dislikes": {
                                "$sum": {
                                    "$cond": [
                                        {"$eq": ["$reaction", "dislike"]},
                                        1,
                                        0,
                                    ]
                                }
                            },
                        }
                    },
                ]
                for row in mongo_collections.document_reactions.aggregate(reaction_pipeline):
                    reaction_map[str(row.get("_id"))] = {
                        "likes": row.get("likes", 0),
                        "dislikes": row.get("dislikes", 0),
                    }
            except Exception as agg_err:
                print("[get_featured_documents_week] aggregate reactions error:", agg_err)
            
            try:
                comment_pipeline = [
                    {"$match": {"documentId": {"$in": doc_ids}}},
                    {"$group": {"_id": "$documentId", "count": {"$sum": 1}}},
                ]
                for row in mongo_collections.document_comments.aggregate(comment_pipeline):
                    comment_map[str(row.get("_id"))] = row.get("count", 0)
            except Exception as agg_err:
                print("[get_featured_documents_week] aggregate comments error:", agg_err)
        
        # Helper function để tính grade
        def calculate_grade(likes, dislikes, comments):
            total = likes + dislikes + comments
            if total == 0:
                return "N/A", "0.0"
            
            # Tính điểm dựa trên tỷ lệ likes và số lượng tương tác
            like_ratio = likes / total if total > 0 else 0
            interaction_score = min(total / 10, 1.0)  # Normalize to 0-1
            
            score = (like_ratio * 0.7 + interaction_score * 0.3) * 10
            
            if score >= 9.0:
                return "A+", f"{score:.1f}"
            elif score >= 8.5:
                return "A", f"{score:.1f}"
            elif score >= 8.0:
                return "A-", f"{score:.1f}"
            elif score >= 7.5:
                return "B+", f"{score:.1f}"
            elif score >= 7.0:
                return "B", f"{score:.1f}"
            elif score >= 6.5:
                return "B-", f"{score:.1f}"
            elif score >= 6.0:
                return "C+", f"{score:.1f}"
            elif score >= 5.5:
                return "C", f"{score:.1f}"
            else:
                return "C-", f"{score:.1f}"
        
        # Format kết quả
        result = []
        for doc in docs:
            doc_id_obj = doc.get("_id")
            doc_id_str = str(doc_id_obj) if doc_id_obj else None
            
            reactions = reaction_map.get(doc_id_str, {})
            likes = reactions.get("likes", 0)
            dislikes = reactions.get("dislikes", 0)
            comment_count = comment_map.get(doc_id_str, 0)
            
            grade, grade_score = calculate_grade(likes, dislikes, comment_count)
            
            # Lấy thông tin school, category, user
            school_name = ""
            category_name = ""
            uploader_name = ""
            
            sid = doc.get("schoolId") or doc.get("school_id")
            if sid:
                try:
                    s = mongo_collections.schools.find_one({"_id": sid if isinstance(sid, ObjectId) else ObjectId(str(sid))})
                    if s:
                        school_name = s.get("name", "")
                except Exception:
                    pass
            
            cid = doc.get("categoryId") or doc.get("category_id")
            if cid:
                try:
                    c = mongo_collections.categories.find_one({"_id": cid if isinstance(cid, ObjectId) else ObjectId(str(cid))})
                    if c:
                        category_name = c.get("name", "")
                except Exception:
                    pass
            
            uid = doc.get("userId") or doc.get("user_id")
            if uid:
                try:
                    u = mongo_collections.users.find_one({"_id": uid if isinstance(uid, ObjectId) else ObjectId(str(uid))})
                    if u:
                        uploader_name = u.get("fullName") or u.get("username") or ""
                except Exception:
                    pass
            
            # Format meta string
            meta_parts = []
            if category_name:
                meta_parts.append(category_name)
            if school_name:
                meta_parts.append(school_name)
            if meta_parts:
                meta = " · ".join(meta_parts)
            else:
                meta = "Tài liệu"
            
            # Lấy keywords làm badges
            keywords = doc.get("keywords", [])
            badges = keywords[:3] if keywords else []
            
            # Format thời gian
            created = doc.get("createdAt") or doc.get("created_at")
            time_ago = ""
            if created:
                if isinstance(created, str):
                    try:
                        created = datetime.fromisoformat(created.replace("Z", "+00:00"))
                    except Exception:
                        created = None
                if created:
                    delta = now - (created.replace(tzinfo=None) if created.tzinfo else created)
                    days = delta.days
                    hours = delta.seconds // 3600
                    minutes = (delta.seconds % 3600) // 60
                    
                    if days > 0:
                        time_ago = f"{days} ngày trước"
                    elif hours > 0:
                        time_ago = f"{hours} giờ trước"
                    elif minutes > 0:
                        time_ago = f"{minutes} phút trước"
                    else:
                        time_ago = "Vừa xong"
            
            # Lấy downloads (nếu có)
            downloads = doc.get("downloads", 0)
            
            result.append({
                "_id": doc_id_str,
                "title": doc.get("title", ""),
                "meta": meta,
                "badges": badges,
                "grade": grade,
                "gradeScore": grade_score,
                "views": doc.get("views", 0),
                "downloads": downloads,
                "time": time_ago,
                "schoolName": school_name,
                "categoryName": category_name,
                "uploaderName": uploader_name,
            })
        
        return jsonify({"documents": result}), 200
        
    except Exception as e:
        print(f"[ERROR] get_featured_documents_week: {e}")
        return jsonify({"error": f"Lỗi server nội bộ: {e}"}), 500


# ===================== DELETE =====================

@documents_bp.route("/<string:doc_id>", methods=["DELETE"])
def delete_document(doc_id):
    """Xóa tài liệu (chỉ user đã upload mới được xóa)."""
    try:
        _id = ObjectId(doc_id)
    except Exception:
        return jsonify({"error": "document id không hợp lệ"}), 400

    try:
        # Xác thực user
        try:
            current_user_oid, _ = _get_current_user_strict()
        except ExpiredSignatureError:
            return jsonify({"error": "JWT hết hạn. Vui lòng đăng nhập lại."}), 401
        except InvalidTokenError as e:
            return jsonify({"error": f"Token không hợp lệ: {e}"}), 401

        # Kiểm tra document có tồn tại và thuộc về user này không
        doc = mongo_collections.documents.find_one({"_id": _id}, {"userId": 1, "user_id": 1})
        if not doc:
            return jsonify({"error": "Không tìm thấy tài liệu"}), 404

        doc_user_id = doc.get("userId") or doc.get("user_id")
        # So sánh ObjectId
        if isinstance(doc_user_id, ObjectId):
            if doc_user_id != current_user_oid:
                return jsonify({"error": "Bạn không có quyền xóa tài liệu này"}), 403
        else:
            if str(doc_user_id) != str(current_user_oid):
                return jsonify({"error": "Bạn không có quyền xóa tài liệu này"}), 403

        # Xóa document
        result = mongo_collections.documents.delete_one({"_id": _id})
        if result.deleted_count == 0:
            return jsonify({"error": "Xóa tài liệu thất bại"}), 500

        # Xóa view history liên quan (optional)
        try:
            mongo_collections.view_history.delete_many({"documentId": _id})
        except Exception:
            pass

        return jsonify({"success": True, "message": "Đã xóa tài liệu"}), 200
    except Exception as e:
        print(f"[ERROR] delete_document: {e}")
        return jsonify({"error": f"Lỗi server nội bộ: {e}"}), 500

